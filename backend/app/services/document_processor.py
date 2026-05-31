from dataclasses import dataclass


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    headings: list[str]
    metadata: dict


class DocumentProcessor:
    def extract(self, filename: str, content: bytes) -> ExtractedDocument:
        suffix = filename.rsplit(".", 1)[-1].lower()
        if suffix == "pdf":
            return self._extract_pdf(content)
        if suffix == "docx":
            return self._extract_docx(content)
        return ExtractedDocument(text=content.decode("utf-8", errors="ignore"), headings=[], metadata={})

    def _extract_pdf(self, content: bytes) -> ExtractedDocument:
        import fitz

        doc = fitz.open(stream=content, filetype="pdf")
        pages = [page.get_text("text") for page in doc]
        text = "\n".join(pages)
        headings = [line.strip() for line in text.splitlines() if line.strip().istitle()][:30]
        return ExtractedDocument(
            text=text,
            headings=headings,
            metadata={"page_count": len(doc), "source_type": "pdf"},
        )

    def _extract_docx(self, content: bytes) -> ExtractedDocument:
        from io import BytesIO

        from docx import Document

        doc = Document(BytesIO(content))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        headings = [
            p.text.strip()
            for p in doc.paragraphs
            if p.style and p.style.name.lower().startswith("heading")
        ]
        return ExtractedDocument(
            text="\n".join(paragraphs),
            headings=headings,
            metadata={"paragraph_count": len(paragraphs), "source_type": "docx"},
        )
